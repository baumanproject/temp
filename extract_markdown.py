import re
from dataclasses import dataclass
from typing import Optional

from docx import Document
from rapidfuzz import fuzz


# -------------------- normalize for matching --------------------

_WS = re.compile(r"\s+")
_JUNK = re.compile(r"[^0-9a-zA-Zа-яА-ЯёЁ%№\s]+")

def norm(s: str) -> str:
    s = (s or "").strip().casefold().replace("ё", "е")
    s = _JUNK.sub(" ", s)
    s = _WS.sub(" ", s).strip()
    return s


# -------------------- row.cells de-dup for merged cells --------------------
# row.cells can return merged cell multiple times -> expected in python-docx  [oai_citation:3‡Stack Overflow](https://stackoverflow.com/questions/48090922/python-docx-row-cells-return-a-merged-cell-multiple-times?utm_source=chatgpt.com)

def unique_row_cells(row):
    out = []
    seen = set()
    for c in row.cells:
        tc = getattr(c, "_tc", None)  # underlying xml cell
        key = id(tc) if tc is not None else id(c)
        if key in seen:
            continue
        seen.add(key)
        out.append(c)
    return out


# -------------------- nested tables as text (ONLY for value) --------------------
# Nested tables are accessible via _Cell.tables / _Cell.iter_inner_content()  [oai_citation:4‡Python-Docx](https://python-docx.readthedocs.io/en/latest/user/tables.html?utm_source=chatgpt.com)

def table_to_text(table, max_depth: int = 2, depth: int = 0) -> str:
    if depth >= max_depth:
        return ""
    lines = []
    for row in table.rows:
        cells = unique_row_cells(row)
        parts = [cell_value_text(c, max_depth=max_depth, depth=depth + 1) for c in cells]
        line = " | ".join([p for p in parts if p.strip()])
        if line.strip():
            lines.append(line.strip())
    return "\n".join(lines).strip()

def cell_value_text(cell, max_depth: int = 2, depth: int = 0) -> str:
    """
    Full text for VALUE:
      - all paragraphs text
      - + nested tables rendered to text (as-is, appended)
    """
    if depth >= max_depth:
        return ""

    parts = []
    for p in cell.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # nested tables are NOT part of cell.text by default; we append them explicitly  [oai_citation:5‡GitHub](https://github.com/python-openxml/python-docx/issues/769?utm_source=chatgpt.com)
    for t in cell.tables:
        tt = table_to_text(t, max_depth=max_depth, depth=depth + 1)
        if tt:
            parts.append(tt)

    return "\n".join(parts).strip()

def cell_key_text(cell, max_lines: int = 2, max_chars: int = 160) -> str:
    """
    Text used ONLY for KEY matching:
      - only paragraphs (no nested tables), to avoid matching on huge nested content
      - take first 1-2 non-empty lines
      - if "Key: Value" -> use left side as key
    """
    lines = []
    for p in cell.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)

    if not lines:
        return ""

    head = " ".join(lines[:max_lines]).strip()
    if len(head) > max_chars:
        head = head[:max_chars].rstrip()

    if ":" in head:
        left = head.split(":", 1)[0].strip()
        if left:
            return left

    return head


# -------------------- output --------------------

@dataclass(frozen=True)
class MatchDetail:
    attribute: str
    found: bool
    score: float
    table_index: Optional[int]
    row_index: Optional[int]
    key_cell_index: Optional[int]


def extract_mapping_docx_top_level(
    docx_path: str,
    attributes: list[str],
    threshold: float = 85.0,
    key_cols: int = 2,
    nested_value_depth: int = 2,
    key_max_lines: int = 2,
    key_max_chars: int = 160,
) -> tuple[dict[str, str | None], float, list[MatchDetail]]:
    """
    - Search: ONLY top-level tables: Document.tables  [oai_citation:6‡Python-Docx](https://python-docx.readthedocs.io/en/latest/api/document.html?utm_source=chatgpt.com)
    - Value: other cells in the same row; nested tables inside value cells appended as text
    """
    doc = Document(docx_path)

    # Pre-collect all rows from top-level tables only
    rows = []  # (ti, ri, key_texts, value_texts)
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            cells = unique_row_cells(row)
            key_texts = [cell_key_text(c, max_lines=key_max_lines, max_chars=key_max_chars) for c in cells]
            value_texts = [cell_value_text(c, max_depth=nested_value_depth) for c in cells]
            rows.append((ti, ri, key_texts, value_texts))

    mapping: dict[str, str | None] = {a: None for a in attributes}
    details: list[MatchDetail] = []
    found_cnt = 0

    for attr in attributes:
        q = norm(attr)

        best_seen = 0.0
        best_meta = (None, None, None)

        matched = False

        # earliest match in document order
        for ti, ri, key_texts, value_texts in rows:
            # match key only among first key_cols cells
            best_score = 0.0
            best_ci = None
            for ci in range(min(key_cols, len(key_texts))):
                kt = key_texts[ci]
                if not kt:
                    continue
                score = float(fuzz.WRatio(q, norm(kt)))
                if score > best_score:
                    best_score = score
                    best_ci = ci

            if best_score > best_seen:
                best_seen = best_score
                best_meta = (ti, ri, best_ci)

            if best_ci is not None and best_score >= threshold:
                # value = all other cells in the same row (same level)
                parts = []
                for ci, txt in enumerate(value_texts):
                    if ci == best_ci:
                        continue
                    t = (txt or "").strip()
                    if t:
                        parts.append(t)

                value = "\n\n".join(parts).strip()

                # If other cells are empty and key cell is "Key: Value", take value from the same cell
                if not value:
                    # use first paragraph line from VALUE text (paragraphs only are fine here too)
                    first_line = ""
                    for p in (value_texts[best_ci] or "").splitlines():
                        if p.strip():
                            first_line = p.strip()
                            break
                    if ":" in first_line:
                        _, right = first_line.split(":", 1)
                        value = right.strip()

                if value:
                    mapping[attr] = value
                    found_cnt += 1
                    details.append(MatchDetail(attr, True, best_score, ti, ri, best_ci))
                else:
                    mapping[attr] = None
                    details.append(MatchDetail(attr, False, best_score, ti, ri, best_ci))

                matched = True
                break

        if not matched:
            ti, ri, ci = best_meta
            details.append(MatchDetail(attr, False, best_seen, ti, ri, ci))

    recall = found_cnt / max(1, len(attributes))
    return mapping, recall, details
