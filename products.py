# pip install python-docx pandas openpyxl
import os
import json
import csv
import re
from collections import Counter

import pandas as pd
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn

# ================== утилиты ================== #

def iter_block_items(document):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)

def _normalize_row_cells(row):
    normalized = []
    prev_xml = None
    for cell in row.cells:
        xml = cell._tc.xml
        if xml == prev_xml:
            continue
        prev_xml = xml
        normalized.append(cell)
    return normalized

# ================== извлечение содержимого ячейки ================== #

def _extract_cell_json(cell):
    texts = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    text = "\n".join(texts) if texts else ""

    nested_tables = list(cell.tables)
    if nested_tables:
        tables_data = []
        for t in nested_tables:
            tables_data.append(_combine_table_block(_table_to_dicts(t)))
        table_field = tables_data
    else:
        table_field = ""

    obj = {"text": text, "table": table_field}
    return json.dumps(obj, ensure_ascii=False)

def _cells_to_inner_table(cells):
    if not cells:
        return []
    max_cols = len(cells)
    headers = [f"col_{i+1}" for i in range(max_cols)]
    row_obj = {}
    for i, cell in enumerate(cells):
        row_obj[headers[i]] = _extract_cell_json(cell)
    return [row_obj]

# ================== определение индекса-колонки ================== #

_index_re = re.compile(r"^\s*\d+([.)])?\s*$")

def _is_index_column(rows, idx, max_rows=5):
    score = 0
    n = min(len(rows), max_rows)
    header = f"col_{idx+1}"
    for row in rows[:n]:
        if header not in row:
            continue
        val_json = row.get(header)
        if not isinstance(val_json, str):
            continue
        try:
            obj = json.loads(val_json)
        except Exception:
            continue
        text = obj.get("text", "").strip()
        table_field = obj.get("table", "")
        if _index_re.match(text) and (table_field == "" or table_field == []):
            score += 1
    return score >= max(2, int(n * 0.6))

# ================== определение основной структуры колонок ================== #

def _determine_main_cols(norm_rows):
    counts = [len(r) for r in norm_rows[:5]]
    if counts and counts[0] == 1:
        counts = counts[1:]
    if not counts:
        return 0, False
    ctr = Counter(counts)
    most_common, freq = ctr.most_common(1)[0]
    if freq > 1:
        cols = most_common
    else:
        cols = min(counts)

    headers = [f"col_{i+1}" for i in range(cols)]
    sample_rows = []
    for r in norm_rows[:5]:
        row_obj = {}
        for i in range(cols):
            if i < len(r):
                row_obj[f"col_{i+1}"] = _extract_cell_json(r[i])
            else:
                row_obj[f"col_{i+1}"] = json.dumps({"text": "", "table": ""}, ensure_ascii=False)
        sample_rows.append(row_obj)

    if _is_index_column(sample_rows, 0):
        return cols - 1, True
    return cols, False

# ================== объединение блоков строк ================== #

def _reduce_and_merge_rows(rows, headers):
    """
    rows: list of dict (col_i -> JSON string)
    headers: list of header names
    Объединяем строки, если первая колонка пустая (""), то склеиваем содержимое остальных колонок
    с предыдущей строкой.
    Возвращает новую list of dict.
    """
    if not rows:
        return rows
    result = []
    prev = rows[0]
    result.append(prev)
    for row in rows[1:]:
        first_cell = json.loads(row[headers[0]])["text"].strip()
        if first_cell == "":
            # склеиваем с prev
            for h in headers[1:]:
                prev_obj = json.loads(prev[h])
                cur_obj = json.loads(row[h])
                # объединяем тексты
                combined_text = (prev_obj["text"] + "\n" + cur_obj["text"]).strip()
                # объединяем tables
                combined_table = prev_obj["table"] if prev_obj["table"] else []
                if cur_obj["table"]:
                    combined_table = combined_table + (cur_obj["table"] if isinstance(cur_obj["table"], list) else [cur_obj["table"]])
                prev[h] = json.dumps({"text": combined_text, "table": combined_table}, ensure_ascii=False)
            continue
        else:
            result.append(row)
            prev = row
    return result

# ================== комбинирование блока таблиц ================== #

def _combine_table_block(list_of_tbl_dicts):
    """
    list_of_tbl_dicts: list of tables, each table represented as list[dict] rows.
    Пытаемся слить их в одну: сначала проверяем одну структуру header & dropped_index.
    Затем объединяем все строки, а потом применяем reduce и удаление индексной колонки.
    Возвращает итог list[dict].
    """
    if not list_of_tbl_dicts:
        return []

    # берем первую таблицу как базу
    base = list_of_tbl_dicts[0]
    headers = list(base[0].keys()) if base else []
    # объединяем все строки
    combined_rows = []
    for tbl in list_of_tbl_dicts:
        combined_rows.extend(tbl)

    # reduce: удаляем индекс колонку, если она оказалось индексной
    if _is_index_column(combined_rows, 0):
        # удаляем header col_1
        new_headers = headers[1:]
        # пересоздаём строки с new_headers
        new_rows = []
        for r in combined_rows:
            new_r = {}
            for idx, h in enumerate(new_headers):
                new_r[f"col_{idx+1}"] = r[h]
            new_rows.append(new_r)
        combined_rows = new_rows
        headers = [f"col_{i+1}" for i in range(len(new_headers))]
    else:
        headers = [f"col_{i+1}" for i in range(len(headers))]

    # reduce и merge пустых первых колонок
    reduced = _reduce_and_merge_rows(combined_rows, headers)
    return reduced

# ================== конвертация таблицы ================== #

def _table_to_dicts(table):
    rows = list(table.rows)
    if not rows:
        return []

    norm_rows = [_normalize_row_cells(r) for r in rows]

    main_cols, dropped_index = _determine_main_cols(norm_rows)
    if main_cols <= 0:
        main_cols = len(norm_rows[0])

    offset = 1 if dropped_index else 0
    headers = [f"col_{i+1}" for i in range(main_cols)]

    result = []
    empty_json = json.dumps({"text": "", "table": ""}, ensure_ascii=False)

    for row_cells in norm_rows:
        row_obj = {h: empty_json for h in headers}
        cell_count = len(row_cells)

        if cell_count - offset <= main_cols:
            for idx in range(main_cols):
                src_idx = idx + offset
                if src_idx < cell_count:
                    row_obj[headers[idx]] = _extract_cell_json(row_cells[src_idx])
                else:
                    row_obj[headers[idx]] = empty_json
        else:
            for idx in range(main_cols):
                src_idx = idx + offset
                if src_idx < cell_count:
                    row_obj[headers[idx]] = _extract_cell_json(row_cells[src_idx])
                else:
                    row_obj[headers[idx]] = empty_json

            extra_cells = row_cells[offset + main_cols:]
            target_idx = None
            for idx in range(main_cols):
                obj = json.loads(row_obj[headers[idx]])
                if obj["text"] == "" and obj["table"] == "":
                    target_idx = idx
                    break
            if target_idx is None:
                target_idx = main_cols - 1

            inner_cells = [row_cells[offset + target_idx]] + extra_cells
            inner_table = _cells_to_inner_table(inner_cells)
            row_obj[headers[target_idx]] = json.dumps(
                {"text": "", "table": inner_table},
                ensure_ascii=False
            )

        result.append(row_obj)

    return result

# ================== работа с docx, поиск блоков и объединение ================== #

def extract_tables_with_text_and_merge(docx_path):
    doc = Document(docx_path)
    blocks = []  # список (text_before, table_rows list[dict])
    text_buffer = []
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            txt = block.text.strip()
            if txt:
                text_buffer.append(txt)
        elif isinstance(block, Table):
            caption = "\n".join(text_buffer).strip() or None
            text_buffer = []
            tbl_dict = _table_to_dicts(block)
            blocks.append((caption, tbl_dict))

    # ищем начала «Паспорт продукта» и конца «Приложение i»
    combined_tables = []
    inside = False
    current_block = []
    for caption, tbl in blocks:
        lower = (caption or "").lower()
        if not inside:
            if "паспорт продукта" in lower:
                inside = True
                current_block = [tbl]
            # иначе пропускаем
        else:
            if re.match(r"приложение\s*\d+", lower):
                # конец блока
                combined_tables.append(current_block)
                inside = False
                current_block = []
            else:
                current_block.append(tbl)
    # если закончились и были открыты
    if inside and current_block:
        combined_tables.append(current_block)

    # теперь склеиваем каждую группу
    merged_results = []
    for tbl_group in combined_tables:
        merged = _combine_table_block(tbl_group)
        merged_results.append(merged)

    # преобразование в нужный формат: список словарей с text_before и df
    results = []
    for merged in merged_results:
        if not merged:
            continue
        df = pd.DataFrame(merged)
        results.append({"text_before": None, "df": df})

    return results

# ================== сохранение / загрузка ================== #

def save_tables_to_csv(tables, base_filename="output"):
    for idx, item in enumerate(tables, start=1):
        fname_csv = f"{base_filename}_table_{idx}.csv"
        fname_xlsx = f"{base_filename}_table_{idx}.xlsx"
        item["df"].to_csv(
            fname_csv,
            index=False,
            encoding="utf-8",
            quoting=csv.QUOTE_ALL
        )
        item["df"].to_excel(
            fname_xlsx,
            index=False,
            engine="openpyxl"
        )

def load_table_from_csv(fname_csv):
    df = pd.read_csv(fname_csv, dtype=str, encoding="utf-8")
    for col in df.columns:
        df[col] = df[col].apply(
            lambda v: json.loads(v) if isinstance(v, str) and v.strip() else {"text":"", "table":""}
        )
    return df

# ================== обход папок ================== #

def process_folder(folder_path):
    results = []
    for fname in os.listdir(folder_path):
        if not fname.lower().endswith(".docx"):
            continue
        full_path = os.path.join(folder_path, fname)
        try:
            tables = extract_tables_with_text_and_merge(full_path)
            serialized = []
            for tbl in tables:
                serialized.append({
                    "table": tbl["df"].to_dict(orient="records")
                })
            results.append({
                "filename": fname,
                "tables": serialized
            })
        except Exception as exc:
            print(f"Ошибка обработки {full_path}: {exc}")
    return results

def main():
    base_dir = "/data/products"
    categories = ["ИСЖ", "ДСЖ", "НСЖ"]
    all_results = {}
    for cat in categories:
        folder = os.path.join(base_dir, cat)
        if not os.path.isdir(folder):
            print(f"Папка не найдена: {folder}")
            continue
        all_results[cat] = process_folder(folder)

    for cat, data in all_results.items():
        out_path = os.path.join(base_dir, f"{cat}_merged_results.json")
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

    return all_results

if __name__ == "__main__":
    results = main()
    for cat, arr in results.items():
        print(f"{cat}: обработано {len(arr)} файлов")
